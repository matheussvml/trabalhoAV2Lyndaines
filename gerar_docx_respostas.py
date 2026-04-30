from pathlib import Path
import re
import shutil
import unicodedata

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt


BASE_DIR = Path(__file__).resolve().parent
MARKDOWN = BASE_DIR / "Respostas.md"
TEMPLATE = BASE_DIR / "Template - Trabalho - 26.1.docx"
DOCX = BASE_DIR / "Template - Trabalho - 26.1 copia.docx"


def split_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(line):
    cells = split_table_row(line)
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


SUBSCRIPTS = str.maketrans("0123456789+-=()", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎")


def simplify_inline_math(text):
    text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    text = re.sub(r"\\hat\{\\beta\}_([0-9])", lambda m: "β̂" + m.group(1).translate(SUBSCRIPTS), text)
    text = re.sub(r"\\beta_([0-9])", lambda m: "β" + m.group(1).translate(SUBSCRIPTS), text)
    text = re.sub(r"X_([0-9])", lambda m: "X" + m.group(1).translate(SUBSCRIPTS), text)
    text = re.sub(r"x_([0-9])", lambda m: "x" + m.group(1).translate(SUBSCRIPTS), text)
    text = re.sub(r"e_([a-zA-Z0-9])", lambda m: "e" + m.group(1).translate(SUBSCRIPTS), text)
    text = text.replace(r"\times", "×")
    text = text.replace(r"\Delta", "Δ")
    text = text.replace(r"\rightarrow", "→")
    text = text.replace(r"\Rightarrow", "⇒")
    text = text.replace(r"\quad", " ")
    text = re.sub(r"\\text\{([^{}]+)\}", r"\1", text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace("\\", "")
    return text


def add_inline_runs(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\$.*?\$)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif part.startswith("$") and part.endswith("$"):
            run = paragraph.add_run(simplify_inline_math(part[1:-1]))
            run.italic = True
        else:
            paragraph.add_run(part)


def add_paragraph_safe(document, style=None):
    try:
        return document.add_paragraph(style=style) if style else document.add_paragraph()
    except KeyError:
        return document.add_paragraph()


def unwrap_latex_command(expr, command):
    target = f"\\{command}" + "{"
    while target in expr:
        start = expr.find(target)
        content_start = start + len(target)
        depth = 1
        pos = content_start
        while pos < len(expr) and depth:
            if expr[pos] == "{":
                depth += 1
            elif expr[pos] == "}":
                depth -= 1
            pos += 1
        if depth:
            break
        inner = expr[content_start : pos - 1]
        expr = expr[:start] + inner + expr[pos:]
    return expr


def add_markdown_table(document, table_lines):
    rows = [split_table_row(line) for line in table_lines if not is_separator_row(line)]
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    for row_idx, row in enumerate(rows):
        for col_idx in range(max_cols):
            text = row[col_idx] if col_idx < len(row) else ""
            cell = table.cell(row_idx, col_idx)
            paragraph = cell.paragraphs[0]
            add_inline_runs(paragraph, text)
            if row_idx == 0:
                for run in paragraph.runs:
                    run.bold = True

    document.add_paragraph()


def add_image(document, alt, image_path):
    full_path = BASE_DIR / image_path
    if not full_path.exists():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[Imagem não encontrada: {image_path}]").italic = True
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(full_path), width=Inches(6.1))

    caption = document.add_paragraph(alt)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.italic = True


def normalize_latex(expr):
    expr = unicodedata.normalize("NFKD", expr).encode("ascii", "ignore").decode("ascii")
    expr = unwrap_latex_command(expr, "boxed")
    expr = expr.replace(r"\checkmark", r"\mathrm{ok}")
    expr = expr.replace(r"\implies", r"\Rightarrow")
    expr = re.sub(r"\\text\{\s*([^{}]+?)\s*\}", lambda m: r"\mathrm{" + m.group(1).replace(" ", r"\;") + "}", expr)
    return expr


def add_equation(document, expr, temp_dir, counter):
    normalized = normalize_latex(expr.strip())
    image_path = temp_dir / f"equacao_{counter:03d}.png"
    width = min(11.5, max(4.5, len(normalized) * 0.085))

    try:
        fig = plt.figure(figsize=(width, 0.65))
        fig.text(0.5, 0.5, f"${normalized}$", ha="center", va="center", fontsize=13)
        plt.axis("off")
        fig.savefig(image_path, dpi=220, bbox_inches="tight", transparent=True, pad_inches=0.08)
        plt.close(fig)

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(min(6.2, width)))
    except Exception:
        plt.close("all")
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(expr)
        run.italic = True


def clear_placeholder(document):
    for paragraph in document.paragraphs:
        if "Início da elaboração" in paragraph.text:
            paragraph.clear()


def fill_student_name(document):
    name = None
    for line in MARKDOWN.read_text(encoding="utf-8").splitlines():
        if line.startswith("**Aluno:**"):
            name = line.replace("**Aluno:**", "").strip()
            break

    if not name:
        return

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == "Digite aqui seu nome":
                    cell.text = name


def build_docx():
    shutil.copyfile(TEMPLATE, DOCX)
    document = Document(str(DOCX))
    clear_placeholder(document)
    fill_student_name(document)

    lines = MARKDOWN.read_text(encoding="utf-8").splitlines()
    table_buffer = []

    temp_dir = BASE_DIR / "_equacoes_docx"
    shutil.rmtree(temp_dir, ignore_errors=True)
    temp_dir.mkdir(exist_ok=True)

    try:
        equation_counter = 1

        for raw_line in lines:
            line = raw_line.rstrip()

            if table_buffer and not line.startswith("|"):
                add_markdown_table(document, table_buffer)
                table_buffer = []

            if not line.strip():
                continue

            if line.startswith("|"):
                table_buffer.append(line)
                continue

            if line.strip() == "---":
                continue

            equation_match = re.fullmatch(r"\$\$(.+)\$\$", line.strip())
            if equation_match:
                add_equation(document, equation_match.group(1), temp_dir, equation_counter)
                equation_counter += 1
                continue

            image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line.strip())
            if image_match:
                add_image(document, image_match.group(1), image_match.group(2))
                continue

            if line.startswith("### "):
                document.add_heading(line[4:].strip(), level=3)
                continue

            if line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
                continue

            if line.startswith("# "):
                title = document.add_heading(line[2:].strip(), level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            if line.startswith("> "):
                paragraph = add_paragraph_safe(document, style="Intense Quote")
                add_inline_runs(paragraph, line[2:].strip())
                continue

            if line.startswith("- "):
                paragraph = add_paragraph_safe(document, style="List Bullet")
                if paragraph.style.name != "List Bullet":
                    paragraph.add_run("• ")
                add_inline_runs(paragraph, line[2:].strip())
                continue

            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, line)

        if table_buffer:
            add_markdown_table(document, table_buffer)

        document.save(str(DOCX))
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


if __name__ == "__main__":
    build_docx()
    print(f"DOCX gerado: {DOCX}")
